import os
import qrcode
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx import Document as DocxDocument


def add_qr_to_docx(docx_path, qr_data, output_path):
    # Создаем QR-код
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=4,
    )
    qr.add_data(qr_data)
    qr.make(fit=True)

    # Генерируем изображение QR-кода
    qr_img = qr.make_image(fill_color="black", back_color="white")
    qr_img_path = "qr_code.png"
    qr_img.save(qr_img_path)

    # Загружаем документ
    doc = DocxDocument(docx_path)

    # Получаем верхний колонтитул (header)
    header = doc.sections[0].header

    # Добавляем изображение QR-кода в колонтитул
    paragraph = header.paragraphs[0]  # Получаем первый параграф в колонтитуле
    run = paragraph.add_run()
    run.add_picture(qr_img_path, width=Inches(0.5), height=Inches(0.5))

    # Выравнивание изображения по правому краю
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Получаем нижний колонтитул (footer)
    footer = doc.sections[0].footer

    # Добавляем изображение QR-кода в нижний колонтитул (левый угол)
    paragraph_footer = footer.paragraphs[0]  # Получаем первый параграф в нижнем колонтитуле
    run_footer = paragraph_footer.add_run()
    run_footer.add_picture(qr_img_path, width=Inches(0.5), height=Inches(0.5))

    # Выравнивание изображения по левому краю
    paragraph_footer.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Сохраняем документ
    doc.save(output_path)

    # Удаляем временный QR-код
    # os.remove(qr_img_path)
# Пример использования
# add_qr_to_docx('input.docx', 'https://example.com', 'output.docx')
